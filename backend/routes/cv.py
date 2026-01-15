"""
CV/Resume routes - Generation, history, download
"""
from fastapi import APIRouter, HTTPException, Depends, Form, File, UploadFile, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from typing import List, Dict, Optional
from datetime import datetime, timezone
import uuid
import re
import json
import logging
from pathlib import Path
from io import BytesIO

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.colors import HexColor, black
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from auth import get_current_user
from database import db
from config import ANTHROPIC_API_KEY, DOWNLOADS_DIR, FREE_LIMITS

# Import shared analysis function for consistent scoring
from routes.resume import analyze_resume_for_role

# Import anthropic
import anthropic
claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY) if ANTHROPIC_API_KEY else None

router = APIRouter(prefix="/cv", tags=["cv"])


# Models
class CVGenerationRequest(BaseModel):
    resume_text: str = Field(..., min_length=100)
    target_role_id: str
    target_region: str = "us"
    region_name: str = "United States"
    target_country: Optional[str] = None
    tier: int = 1
    experience_level: str = "mid"
    experience_years: int = 3
    recommended_length: str = "1 page"
    current_role: Optional[str] = None
    years_experience: Optional[int] = None
    skills: List[str] = []


class CVDownloadRequest(BaseModel):
    cv_type: str = "hybrid"
    format: str = "pdf"


def get_cv_generation_prompt():
    """System prompt for SUPERIOR hybrid resume generation - Best in market"""
    return """You are the world's #1 AI resume writer. Your resumes have a 95%+ interview callback rate.

You create ONE SUPERIOR HYBRID RESUME that:
✅ PASSES EVERY ATS (Applicant Tracking System) - 98%+ ATS compatibility
✅ IMPRESSES HUMAN RECRUITERS - Engaging, achievement-focused storytelling
✅ PERFECTLY OPTIMIZED for the target AI/ML role

=== ATS MASTERY (Your resume MUST pass these) ===

1. KEYWORD OPTIMIZATION (Critical for ATS)
   - Extract ALL keywords from target role requirements
   - Include exact job title variations (ML Engineer, Machine Learning Engineer, MLE)
   - Use BOTH spelled-out AND abbreviated terms (Natural Language Processing, NLP)
   - Place critical keywords in: Header, Skills, First bullet of each job
   - Target: 40-60 relevant keywords naturally integrated

2. ATS-SAFE FORMATTING
   - NO tables, columns, graphics, icons, or text boxes
   - NO headers/footers (ATS can't read them)
   - Standard section headings: SKILLS, EXPERIENCE, EDUCATION, CERTIFICATIONS
   - Dates in format: Month Year - Month Year (e.g., Jan 2022 - Present)
   - Simple bullet points (use • or -)
   - Standard fonts assumed (parsed as plain text)
   - NO special characters except: • - / ( ) , . : |

3. CONTACT SECTION (Top, single line or two lines max)
   - Full Name
   - Email | Phone | LinkedIn URL | GitHub URL | Location (City, State)

=== HUMAN APPEAL (What makes recruiters excited) ===

4. POWERFUL BULLETS (Each one must have ALL THREE)
   - ACTION VERB (Led, Built, Deployed, Optimized, Reduced, Increased, Designed)
   - TECHNOLOGY/SKILL (Python, PyTorch, AWS, Kubernetes, specific tools)
   - QUANTIFIED RESULT (%, $, time saved, scale, users impacted)
   
   GOOD: "Deployed production ML pipeline using PyTorch and AWS SageMaker, reducing inference latency by 40% and serving 2M+ daily predictions"
   BAD: "Worked on machine learning projects using Python"

5. ACHIEVEMENT DENSITY
   - Every bullet = a mini success story
   - Numbers in 80%+ of bullets
   - Impact on business/team/product, not just tasks
   - Show SCOPE (team size, data size, user base, revenue impact)

6. SKILLS SECTION (ATS goldmine + human scannable)
   Format as categories:
   - Languages: Python, SQL, Java, Scala
   - ML/AI: PyTorch, TensorFlow, Scikit-learn, XGBoost, LLMs, RAG, Transformers
   - Cloud & MLOps: AWS (SageMaker, EC2, S3), GCP, Docker, Kubernetes, MLflow
   - Data: Spark, Pandas, PostgreSQL, MongoDB, Redis
   - Tools: Git, Linux, Airflow, Weights & Biases

=== OUTPUT STRUCTURE ===

1. HEADER (Name + Contact)
2. SKILLS (Categorized, keyword-rich - 4-6 lines)
3. EXPERIENCE (2-3 positions, 4-5 bullets each, reverse chronological)
4. PROJECTS (Optional - 1-2 significant projects if space allows)
5. EDUCATION (1-2 lines max)
6. CERTIFICATIONS (Only relevant ones, 1-2 lines)

=== STRICT RULES ===

- EXACTLY 1 PAGE (400-500 words max)
- NO professional summary/objective (wastes space, ATS doesn't care)
- NO "References available upon request"
- NO personal pronouns (I, my, me)
- NO generic phrases ("responsible for", "worked on", "helped with")
- PLAIN TEXT ONLY - No markdown, no special formatting
- Prioritize RECENT experience (last 5-7 years)
- Tailor EVERYTHING to the target role

=== OUTPUT FORMAT (JSON) ===

{
    "resume": {
        "content": "Full resume text here with proper line breaks...",
        "ats_score": 96,
        "human_appeal_score": 94,
        "word_count": 450,
        "keyword_count": 52,
        "keywords_used": ["Python", "PyTorch", "Machine Learning", "AWS", "..."],
        "metrics_count": 12,
        "action_verbs_used": ["Deployed", "Optimized", "Led", "Built", "..."]
    },
    "analysis": {
        "target_role": "Machine Learning Engineer",
        "match_score": 92,
        "strengths": ["Strong ML fundamentals", "Production experience", "..."],
        "improvements_made": ["Added metrics to all bullets", "Optimized keyword placement", "..."],
        "skills_highlighted": ["PyTorch", "AWS", "MLOps", "..."],
        "skills_gap": ["Kubernetes", "Spark"] 
    },
    "ats_breakdown": {
        "keyword_optimization": 95,
        "formatting_compliance": 100,
        "section_structure": 98,
        "contact_info": 100,
        "date_formatting": 100,
        "overall_ats_pass_rate": "98%"
    }
}

Generate the BEST resume that will get this candidate interviews at top AI companies."""


def get_region_standards(region: str, exp_level: str, tier: int) -> str:
    """Get region-specific resume standards"""
    standards = {
        "us": "1 page max, no photo, metrics-focused, action verbs",
        "uk": "1-2 pages, no photo, achievement-focused",
        "eu": "1-2 pages, may include photo, skills matrix common",
        "india": "2-3 pages acceptable, detailed project descriptions",
        "global": "1-2 pages, internationally recognized format"
    }
    return standards.get(region, standards["us"])


def generate_cv_pdf(cv_data: Dict, cv_type: str, target_role: str, user_name: str = "User") -> str:
    """Generate PDF from CV data"""
    file_id = str(uuid.uuid4())[:8]
    filename = f"cv_{cv_type}_{file_id}.pdf"
    filepath = DOWNLOADS_DIR / filename
    
    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=letter,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    styles = getSampleStyleSheet()
    
    name_style = ParagraphStyle(
        'Name',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=6,
        alignment=TA_CENTER
    )
    
    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading2'],
        fontSize=11,
        spaceBefore=10,
        spaceAfter=4,
        textColor=HexColor('#1a365d')
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        leading=12,
        spaceAfter=3
    )
    
    story = []
    
    content = cv_data.get("content", "")
    lines = content.split('\n')
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
        elif i == 0 or line.isupper():
            story.append(Paragraph(line, name_style if i == 0 else section_style))
        else:
            # Clean up bullets
            if line.startswith('•') or line.startswith('-'):
                line = '• ' + line.lstrip('•-').strip()
            story.append(Paragraph(line, body_style))
    
    doc.build(story)
    return str(filepath)


def generate_cv_docx(cv_data: Dict, cv_type: str, target_role: str, user_name: str = "User") -> str:
    """Generate DOCX from CV data"""
    file_id = str(uuid.uuid4())[:8]
    filename = f"cv_{cv_type}_{file_id}.docx"
    filepath = DOWNLOADS_DIR / filename
    
    doc = Document()
    
    content = cv_data.get("content", "")
    lines = content.split('\n')
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        elif i == 0:
            para = doc.add_heading(line, level=0)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.isupper():
            doc.add_heading(line, level=1)
        else:
            doc.add_paragraph(line)
    
    doc.save(str(filepath))
    return str(filepath)


@router.post("/generate")
async def generate_cv_standalone(
    request: CVGenerationRequest,
    user: dict = Depends(get_current_user)
):
    """Generate SUPERIOR hybrid resume - ATS-optimized + Human-appealing"""
    # Import AI_ROLES from server
    from server import AI_ROLES
    
    current_month = datetime.now(timezone.utc)
    usage = await db.usage.find_one({
        "user_id": user["id"],
        "month": current_month.month,
        "year": current_month.year
    }, {"_id": 0})
    
    if not usage:
        usage = {"cv_generations_used": 0}
        await db.usage.insert_one({
            "user_id": user["id"],
            "month": current_month.month,
            "year": current_month.year,
            "cv_generations_used": 0
        })
    
    cv_used = usage.get("cv_generations_used", 0)
    cv_credits = user.get("cv_credits", 0)
    is_pro = user.get("subscription_tier") == "pro"
    
    if not is_pro and cv_used >= FREE_LIMITS.get("cv_generations", 1) and cv_credits <= 0:
        raise HTTPException(
            status_code=403,
            detail={
                "message": "CV generation limit reached",
                "used": cv_used,
                "limit": FREE_LIMITS.get("cv_generations", 1),
                "credits": cv_credits
            }
        )
    
    target_role = next((r for r in AI_ROLES if r["id"] == request.target_role_id), None)
    if not target_role:
        raise HTTPException(status_code=404, detail="Target role not found")
    
    region_standards = get_region_standards(request.target_region, request.experience_level, request.tier)
    
    # Build comprehensive user message for SUPERIOR resume
    user_message = f"""Create the ULTIMATE HYBRID RESUME for this candidate targeting: {target_role['name']}

=== TARGET ROLE REQUIREMENTS ===
Role: {target_role['name']}
Required Skills: {', '.join(target_role.get('top_skills', []))}
Experience Level: {request.experience_years} years ({request.experience_level})
Target Region: {request.target_country or request.region_name}
Regional Standards: {region_standards}

=== CANDIDATE'S CURRENT RESUME ===
{request.resume_text[:5000]}

=== YOUR TASK ===
Transform this resume into a SUPERIOR hybrid version that:
1. Scores 95%+ on ATS systems (Workday, Greenhouse, Lever, Taleo)
2. Impresses hiring managers at {target_role['name']} roles
3. Highlights relevant skills: {', '.join(target_role.get('top_skills', [])[:8])}
4. Includes 40-60 relevant keywords naturally
5. Has metrics/numbers in 80%+ of bullets
6. Fits on exactly 1 page

Make this the BEST resume this candidate has ever had."""
    
    try:
        response = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2500,  # Reduced from 4000 - only 1 version needed
            system=get_cv_generation_prompt(),
            messages=[{"role": "user", "content": user_message}]
        )
        
        response_text = response.content[0].text
        response_text = re.sub(r'```json\s*', '', response_text)
        response_text = re.sub(r'```\s*', '', response_text)
        
        cv_data = json.loads(response_text)
        
        # Transform to versions array for backward compatibility
        if "resume" in cv_data:
            resume = cv_data["resume"]
            analysis = cv_data.get("analysis", {})
            ats_breakdown = cv_data.get("ats_breakdown", {})
            
            cv_data["versions"] = [{
                "type": "hybrid",
                "name": "Superior Hybrid Resume",
                "content": resume.get("content", ""),
                "ats_score": resume.get("ats_score", 95),
                "human_appeal_score": resume.get("human_appeal_score", 94),
                "word_count": resume.get("word_count", 450),
                "keyword_count": resume.get("keyword_count", 50),
                "keywords_used": resume.get("keywords_used", []),
                "metrics_count": resume.get("metrics_count", 10),
                "action_verbs_used": resume.get("action_verbs_used", [])
            }]
            cv_data["key_improvements"] = analysis.get("improvements_made", [])
            cv_data["missing_skills"] = analysis.get("skills_gap", [])
            cv_data["strengths"] = analysis.get("strengths", [])
            cv_data["match_score"] = analysis.get("match_score", 90)
            cv_data["ats_breakdown"] = ats_breakdown
            
    except json.JSONDecodeError as e:
        logging.error(f"CV JSON error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse CV response")
    except Exception as e:
        logging.error(f"CV generation error: {e}")
        raise HTTPException(status_code=500, detail=f"CV generation failed: {str(e)}")
    
    # Update usage
    if not is_pro:
        if cv_credits > 0:
            await db.users.update_one(
                {"id": user["id"]},
                {"$inc": {"cv_credits": -1}}
            )
        else:
            await db.usage.update_one(
                {"user_id": user["id"], "month": current_month.month, "year": current_month.year},
                {"$inc": {"cv_generations_used": 1}},
                upsert=True
            )
    
    # Extract the hybrid resume content for backward compatibility
    hybrid_version = cv_data.get("versions", [{}])[0] if cv_data.get("versions") else {}
    hybrid_content = hybrid_version.get("content", "")
    
    # IMPORTANT: Run the same analysis used by Scanner for consistent, REAL scores
    # This replaces the self-proclaimed scores with verified analysis
    try:
        verified_analysis = await analyze_resume_for_role(hybrid_content, request.target_role_id)
        
        # Update the version with VERIFIED scores (not self-proclaimed)
        if cv_data.get("versions"):
            cv_data["versions"][0]["ats_score"] = verified_analysis.get("ats_score", 85)
            cv_data["versions"][0]["human_appeal_score"] = verified_analysis.get("human_appeal_score", 80)
            cv_data["versions"][0]["keywords_used"] = verified_analysis.get("keywords_found", [])
            
        # Update cv_data with verified analysis
        cv_data["verified_analysis"] = verified_analysis
        hybrid_version = cv_data.get("versions", [{}])[0]
        
        logging.info(f"CV verified: ATS={verified_analysis.get('ats_score')}, Human={verified_analysis.get('human_appeal_score')}")
    except Exception as e:
        logging.warning(f"CV verification failed, using generated scores: {e}")
    
    cv_id = str(uuid.uuid4())
    
    # Store with both new and legacy formats
    cv_record = {
        "id": cv_id,
        "user_id": user["id"],
        "target_role": target_role["name"],
        "target_role_id": request.target_role_id,
        "target_region": request.target_region,
        "versions": cv_data.get("versions", []),
        # Legacy fields for backward compatibility
        "natural_cv": hybrid_content,
        "ats_cv": hybrid_content,  # Same content - it's hybrid
        "natural_resume": hybrid_content,
        "ats_resume": hybrid_content,
        # New enhanced fields
        "key_improvements": cv_data.get("key_improvements", []),
        "missing_skills": cv_data.get("missing_skills", []),
        "strengths": cv_data.get("strengths", []),
        "match_score": cv_data.get("match_score", 0),
        "ats_breakdown": cv_data.get("ats_breakdown", {}),
        "ats_score_estimate": hybrid_version.get("ats_score", 95),
        "human_voice_score": hybrid_version.get("human_appeal_score", 94) / 10,
        "keywords_added": hybrid_version.get("keywords_used", []),
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.cv_generations.insert_one(cv_record)
    
    # Return with both new and legacy formats
    return {
        "cv_id": cv_id,
        "target_role": target_role["name"],
        "versions": cv_data.get("versions", []),
        # Legacy fields for frontend compatibility
        "natural_cv": hybrid_content,
        "ats_cv": hybrid_content,
        "natural_resume": hybrid_content,
        "ats_resume": hybrid_content,
        # Enhanced fields
        "key_improvements": cv_data.get("key_improvements", []),
        "missing_skills": cv_data.get("missing_skills", []),
        "strengths": cv_data.get("strengths", []),
        "match_score": cv_data.get("match_score", 0),
        "ats_breakdown": cv_data.get("ats_breakdown", {}),
        "ats_score_estimate": hybrid_version.get("ats_score", 95),
        "human_voice_score": hybrid_version.get("human_appeal_score", 94) / 10,
        "keywords_added": hybrid_version.get("keywords_used", []),
        # Include verified analysis for consistent scoring transparency
        "verified_analysis": cv_data.get("verified_analysis", {}),
        "usage": {
            "used": cv_used + 1 if not is_pro and cv_credits <= 0 else cv_used,
            "limit": 999 if is_pro else FREE_LIMITS.get("cv_generations", 1),
            "credits_remaining": max(0, cv_credits - 1) if cv_credits > 0 and not is_pro else cv_credits
        }
    }


@router.get("/history")
async def get_cv_history(user: dict = Depends(get_current_user)):
    """Get user's CV generation history"""
    cvs = await db.cv_generations.find(
        {"user_id": user["id"]},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    return {"cv_generations": cvs}


@router.get("/{cv_id}")
async def get_cv(cv_id: str, user: dict = Depends(get_current_user)):
    """Get specific CV generation"""
    cv = await db.cv_generations.find_one(
        {"id": cv_id, "user_id": user["id"]},
        {"_id": 0}
    )
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")
    return cv


@router.post("/{cv_id}/download")
async def download_cv(
    cv_id: str,
    cv_type: str = Form("hybrid"),
    format: str = Form("pdf"),
    user: dict = Depends(get_current_user)
):
    """Download CV as PDF or DOCX"""
    cv = await db.cv_generations.find_one(
        {"id": cv_id, "user_id": user["id"]},
        {"_id": 0}
    )
    
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")
    
    versions = cv.get("versions", [])
    cv_data = next((v for v in versions if v.get("type") == cv_type), versions[0] if versions else None)
    
    if not cv_data:
        raise HTTPException(status_code=404, detail="CV version not found")
    
    target_role = cv.get("target_role", "AI Role")
    user_name = user.get("name", "User")
    
    try:
        if format == "docx":
            filepath = generate_cv_docx(cv_data, cv_type, target_role, user_name)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            filepath = generate_cv_pdf(cv_data, cv_type, target_role, user_name)
            media_type = "application/pdf"
        
        filename = Path(filepath).name
        
        return FileResponse(
            path=filepath,
            filename=filename,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        logging.error(f"CV download error: {e}")
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")


class CVDownloadRequest(BaseModel):
    cv_content: str
    user_name: Optional[str] = "User"


@router.post("/download-direct")
async def download_cv_direct(
    request: CVDownloadRequest,
    cv_version: str = Query("hybrid"),
    format: str = Query("pdf"),
    target_role: str = Query("AI Role"),
    user: dict = Depends(get_current_user)
):
    """Download CV directly from content without saving"""
    cv_data = {"content": request.cv_content, "type": cv_version}
    user_name = request.user_name or user.get("name", "User")
    
    try:
        if format == "docx":
            filepath = generate_cv_docx(cv_data, cv_version, target_role, user_name)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            filepath = generate_cv_pdf(cv_data, cv_version, target_role, user_name)
            media_type = "application/pdf"
        
        filename = Path(filepath).name
        
        return FileResponse(
            path=filepath,
            filename=filename,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        logging.error(f"Direct CV download error: {e}")
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")
