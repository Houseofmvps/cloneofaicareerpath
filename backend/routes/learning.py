"""
Learning Path routes - Generation, progress tracking, downloads
"""
from fastapi import APIRouter, HTTPException, Depends, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from typing import List, Dict, Optional
from datetime import datetime, timezone
import uuid
import re
import json
import logging
from pathlib import Path
from io import BytesIO

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from auth import get_current_user
from database import db
from config import ANTHROPIC_API_KEY, DOWNLOADS_DIR, FREE_LIMITS

# Import anthropic
import anthropic
claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY) if ANTHROPIC_API_KEY else None

router = APIRouter(prefix="/learning-path", tags=["learning"])


# Models
class LearningPathRequest(BaseModel):
    # Support both target_role (string name) and target_role_id (role ID)
    target_role: Optional[str] = None
    target_role_id: Optional[str] = None
    current_role: Optional[str] = None
    years_experience: Optional[int] = 0
    current_skills: List[str] = []
    experience_level: str = "beginner"
    available_hours_per_week: int = 10
    preferred_learning_style: str = "video"
    budget: str = "free"
    location: Optional[str] = None


class WeekProgressUpdate(BaseModel):
    week: int
    completed: bool
    notes: Optional[str] = None


class CourseProgressUpdate(BaseModel):
    week: int
    course_index: int
    completed: bool
    notes: Optional[str] = None


class SaveCourseRequest(BaseModel):
    course_name: str
    course_url: str
    platform: str
    week: Optional[int] = None
    learning_path_id: Optional[str] = None
    notes: Optional[str] = None


def get_learning_path_prompt():
    """System prompt for learning path generation - kept for backward compatibility"""
    return """You are an expert AI/ML career coach."""


def build_verified_learning_path(
    role_path: Optional[Dict],
    target_role_name: str,
    target_role_id: str,
    current_skills: List[str],
    experience_level: str,
    hours_per_week: int,
    budget: str,
    current_role: str = None
) -> Dict:
    """
    Build a learning path using VERIFIED courses from our database.
    No AI hallucinations - only real courses with working URLs.
    Each week has ONE focused course - no repetition.
    """
    from data.courses_database import ALL_COURSES, get_course_by_id
    
    # Filter courses by budget preference
    def course_matches_budget(course, budget):
        if budget == "free":
            return course.get("cost_type") in ["free", "freemium"]
        return True
    
    # Build weeks from role path
    weeks = []
    
    if role_path and role_path.get("weekly_curriculum"):
        for week_data in role_path["weekly_curriculum"]:
            week_num = week_data["week"]
            theme = week_data["theme"]
            focus = week_data["focus"]
            course_ids = week_data.get("courses", [])
            
            # Get actual course objects
            week_courses = []
            for course_id in course_ids:
                course = get_course_by_id(course_id)
                if course:
                    # For free budget, mark paid courses as "optional"
                    is_optional = budget == "free" and course.get("cost_type") == "paid"
                    week_courses.append({
                        "id": course.get("id"),
                        "name": course["name"],
                        "platform": course["platform"],
                        "url": course["url"],
                        "duration_hours": course.get("duration_hours", 5),
                        "cost": course.get("cost", "Free"),
                        "cost_type": course.get("cost_type", "free"),
                        "badge": course.get("badge", ""),
                        "rating": course.get("rating", 4.5),
                        "instructor": course.get("instructor", ""),
                        "description": course.get("description", ""),
                        "skills_taught": course.get("skills_taught", []),
                        "why_recommended": course.get("why_recommended", ""),
                        "is_optional": is_optional
                    })
            
            # Add week entry
            weeks.append({
                "week": week_num,
                "theme": theme,
                "focus": focus,
                "hours": hours_per_week,
                "courses": week_courses,
                "project": {
                    "name": f"Week {week_num}: {theme} Project",
                    "description": f"Apply {theme.lower()} concepts in a hands-on project",
                    "skills_practiced": week_courses[0].get("skills_taught", ["Python"])[:3] if week_courses else ["Python"]
                },
                "milestone": focus
            })
    else:
        # Fallback for unknown roles - generic path
        fallback_sequence = [
            (1, "Python Foundations", "py-001", "Master Python basics"),
            (2, "Math for ML", "math-002", "Linear algebra intuition"),
            (3, "Statistics", "stats-001", "Statistical foundations"),
            (4, "ML Quick Start", "ml-003", "Google ML Crash Course"),
            (5, "ML Foundations", "ml-001", "Andrew Ng's ML course"),
            (6, "Practical ML", "ml-002", "fast.ai approach"),
            (7, "Deep Learning", "dl-001", "Neural networks"),
            (8, "PyTorch", "dl-002", "PyTorch framework"),
            (9, "AI Engineering", "ai-001", "Scrimba AI Engineer"),
            (10, "LangChain", "ai-002", "Build with LangChain"),
            (11, "MLOps", "mlops-002", "Made With ML"),
            (12, "Interview Prep", "int-001", "Ace the interview"),
        ]
        
        for week_num, theme, course_id, focus in fallback_sequence:
            course = get_course_by_id(course_id)
            if course and course_matches_budget(course, budget):
                weeks.append({
                    "week": week_num,
                    "theme": theme,
                    "focus": focus,
                    "hours": hours_per_week,
                    "courses": [{
                        "id": course.get("id"),
                        "name": course["name"],
                        "platform": course["platform"],
                        "url": course["url"],
                        "duration_hours": course.get("duration_hours", 5),
                        "cost": course.get("cost", "Free"),
                        "cost_type": course.get("cost_type", "free"),
                        "badge": course.get("badge", ""),
                        "rating": course.get("rating", 4.5),
                        "description": course.get("description", ""),
                        "skills_taught": course.get("skills_taught", [])
                    }],
                    "project": {
                        "name": f"Week {week_num} Project",
                        "description": f"Apply {theme.lower()} concepts",
                        "skills_practiced": ["Python", "ML"]
                    },
                    "milestone": focus
                })
    
    # Build complete path data
    path_data = {
        "path_overview": {
            "target_role": target_role_name,
            "target_role_id": target_role_id,
            "duration_weeks": len(weeks),
            "hours_per_week": hours_per_week,
            "total_hours": len(weeks) * hours_per_week,
            "experience_level": experience_level,
            "current_role": current_role,
            "budget_preference": budget,
            "difficulty_progression": "Beginner → Intermediate → Advanced → Production"
        },
        "weeks": weeks,
        "fast_track": {
            "enabled": True,
            "title": "⚡ Fast Track Your Learning",
            "description": "Complete your journey 3x faster with interactive, hands-on courses",
            "courses": [
                {
                    "name": "The AI Engineer Path",
                    "platform": "Scrimba",
                    "url": "https://scrimba.com/the-ai-engineer-path-c02v?via=u436b310",
                    "description": "Complete AI curriculum - OpenAI, LangChain, RAG, Agents",
                    "duration": "40 hours",
                    "why": "Interactive coding. Build real AI apps while learning.",
                    "badge": "🔥 #1 AI Course"
                },
                {
                    "name": "Learn Python",
                    "platform": "Scrimba",
                    "url": "https://scrimba.com/learn-python-c03?via=u436b310",
                    "description": "Python from scratch with interactive exercises",
                    "duration": "15 hours",
                    "why": "Code in your browser. No setup needed.",
                    "badge": "🎯 Interactive"
                }
            ],
            "benefits": [
                "✅ Interactive coding in browser",
                "✅ Build real projects",
                "✅ AI-focused curriculum",
                "✅ Up-to-date content (2024-2025)"
            ]
        },
        "career_readiness_checklist": [
            "✅ Complete all course projects",
            "✅ Build 3-5 portfolio projects on GitHub",
            "✅ Achieve top 20% in a Kaggle competition",
            "✅ Update LinkedIn with new skills",
            "✅ Practice 50+ interview questions",
            "✅ Network with ML engineers on LinkedIn/Twitter"
        ],
        "recommended_tools": [
            "VS Code with Python extensions",
            "Jupyter Notebooks",
            "Git & GitHub",
            "Google Colab (free GPU)",
            "Weights & Biases (experiment tracking)"
        ],
        "communities_to_join": [
            {"name": "r/MachineLearning", "url": "https://reddit.com/r/MachineLearning"},
            {"name": "MLOps Community", "url": "https://mlops.community/"},
            {"name": "Hugging Face Discord", "url": "https://huggingface.co/join/discord"},
            {"name": "fast.ai Forums", "url": "https://forums.fast.ai/"}
        ],
        "interview_prep": {
            "technical": ["ML algorithms", "System design", "Coding challenges"],
            "behavioral": ["STAR method", "Project walkthroughs"],
            "resources": [
                {"name": "Ace the Data Science Interview", "url": "https://www.acethedatascienceinterview.com/"},
                {"name": "ML System Design Primer", "url": "https://github.com/chiphuyen/machine-learning-systems-design"}
            ]
        }
    }
    
    return path_data


@router.post("/generate")
async def generate_learning_path_standalone(
    request: LearningPathRequest,
    user: dict = Depends(get_current_user)
):
    """Standalone Learning Path Generator - generates personalized plan with VERIFIED courses"""
    
    # Import verified course database
    from data.roles import AI_ROLES
    from data.courses_database import (
        ROLE_LEARNING_PATHS, ALL_COURSES, get_course_by_id, get_role_path
    )
    
    # Resolve target_role from target_role_id
    target_role_id = request.target_role_id
    target_role_name = request.target_role
    
    if not target_role_name and target_role_id:
        role = next((r for r in AI_ROLES if r["id"] == target_role_id), None)
        if role:
            target_role_name = role["name"]
        else:
            target_role_name = target_role_id.replace("_", " ").title()
    
    if not target_role_name:
        raise HTTPException(status_code=400, detail="Please specify a target role")
    
    # Determine experience level from years
    experience_level = request.experience_level
    if request.years_experience:
        if request.years_experience < 2:
            experience_level = "beginner"
        elif request.years_experience < 5:
            experience_level = "intermediate"
        else:
            experience_level = "advanced"
    
    # Check usage limits
    current_month = datetime.now(timezone.utc)
    usage = await db.usage.find_one({
        "user_id": user["id"],
        "month": current_month.month,
        "year": current_month.year
    }, {"_id": 0})
    
    if not usage:
        usage = {"learning_paths_used": 0}
        await db.usage.update_one(
            {"user_id": user["id"], "month": current_month.month, "year": current_month.year},
            {"$set": {"learning_paths_used": 0}},
            upsert=True
        )
    
    learning_paths_used = usage.get("learning_paths_used", 0)
    is_pro = user.get("subscription_tier") == "pro"
    free_limit = FREE_LIMITS.get("learning_paths", 1)
    
    if not is_pro and learning_paths_used >= free_limit:
        raise HTTPException(
            status_code=403,
            detail={
                "message": "Learning path limit reached",
                "used": learning_paths_used,
                "limit": free_limit,
                "upgrade_message": "Upgrade to Pro for unlimited learning paths"
            }
        )
    
    # Get role-specific learning path from verified database
    role_path = get_role_path(target_role_id) if target_role_id else None
    
    # Build learning path from verified courses
    path_data = build_verified_learning_path(
        role_path=role_path,
        target_role_name=target_role_name,
        target_role_id=target_role_id,
        current_skills=request.current_skills,
        experience_level=experience_level,
        hours_per_week=request.available_hours_per_week,
        budget=request.budget,
        current_role=request.current_role
    )
    
    # Update usage
    if not is_pro:
        await db.usage.update_one(
            {"user_id": user["id"], "month": current_month.month, "year": current_month.year},
            {"$inc": {"learning_paths_used": 1}},
            upsert=True
        )
    
    # Save to database
    path_id = str(uuid.uuid4())
    await db.learning_paths.insert_one({
        "id": path_id,
        "user_id": user["id"],
        "target_role": target_role_name,
        "target_role_id": request.target_role_id,
        "experience_level": experience_level,
        "current_role": request.current_role,
        "current_skills": request.current_skills,
        "available_hours": request.available_hours_per_week,
        "location": request.location,
        "path_data": path_data,
        "overall_progress": 0,
        "week_progress": {},
        "course_progress": {},
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    return {
        "id": path_id,
        "path_id": path_id,
        "target_role": target_role_name,
        "target_role_id": request.target_role_id,
        "learning_path": path_data,
        "path_data": path_data,
        "usage": {
            "used": learning_paths_used + 1 if not is_pro else learning_paths_used,
            "limit": 999 if is_pro else free_limit
        }
    }


@router.get("/history")
async def get_learning_path_history(user: dict = Depends(get_current_user)):
    """Get user's learning path history"""
    paths = await db.learning_paths.find(
        {"user_id": user["id"]},
        {"_id": 0}
    ).sort("created_at", -1).to_list(20)
    return {"learning_paths": paths}


@router.get("/all-progress")
async def get_all_learning_progress(user: dict = Depends(get_current_user)):
    """Get progress for all learning paths"""
    paths = await db.learning_paths.find(
        {"user_id": user["id"]},
        {"_id": 0, "id": 1, "target_role": 1, "overall_progress": 1, "created_at": 1}
    ).sort("created_at", -1).to_list(20)
    
    saved_courses = await db.saved_courses.find(
        {"user_id": user["id"]},
        {"_id": 0}
    ).to_list(100)
    
    return {
        "learning_paths": paths,
        "saved_courses": saved_courses,
        "total_saved_courses": len(saved_courses)
    }


@router.get("/{path_id}")
async def get_learning_path(path_id: str, user: dict = Depends(get_current_user)):
    """Get specific learning path"""
    path = await db.learning_paths.find_one(
        {"id": path_id, "user_id": user["id"]},
        {"_id": 0}
    )
    if not path:
        raise HTTPException(status_code=404, detail="Learning path not found")
    return path


@router.post("/{path_id}/download")
async def download_learning_path(
    path_id: str,
    format: str = Form("pdf"),
    user: dict = Depends(get_current_user)
):
    """Download learning path as PDF or DOCX"""
    path = await db.learning_paths.find_one(
        {"id": path_id, "user_id": user["id"]},
        {"_id": 0}
    )
    
    if not path:
        raise HTTPException(status_code=404, detail="Learning path not found")
    
    path_data = path.get("path_data", {})
    target_role = path.get("target_role", "AI Role")
    file_id = str(uuid.uuid4())[:8]
    
    if format == "docx":
        filename = f"learning_path_{file_id}.docx"
        filepath = DOWNLOADS_DIR / filename
        
        doc = Document()
        doc.add_heading(f"Learning Path: {target_role}", 0)
        
        overview = path_data.get("path_overview", {})
        doc.add_paragraph(f"Duration: {overview.get('duration_weeks', 16)} weeks")
        doc.add_paragraph(f"Hours/week: {overview.get('hours_per_week', 10)}")
        
        for week in path_data.get("weeks", []):
            doc.add_heading(f"Week {week.get('week')}: {week.get('focus', '')}", level=1)
            doc.add_paragraph(f"Phase: {week.get('phase', '')}")
            
            for course in week.get("courses", []):
                doc.add_paragraph(f"• {course.get('name')} ({course.get('platform')})")
        
        doc.save(str(filepath))
    else:
        filename = f"learning_path_{file_id}.pdf"
        filepath = DOWNLOADS_DIR / filename
        
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=20)
        story.append(Paragraph(f"Learning Path: {target_role}", title_style))
        story.append(Spacer(1, 12))
        
        for week in path_data.get("weeks", []):
            story.append(Paragraph(f"Week {week.get('week')}: {week.get('focus', '')}", styles['Heading2']))
            for course in week.get("courses", []):
                story.append(Paragraph(f"• {course.get('name')} - {course.get('platform')}", styles['Normal']))
            story.append(Spacer(1, 8))
        
        doc.build(story)
    
    return FileResponse(
        path=str(filepath),
        filename=filename,
        media_type="application/octet-stream"
    )


@router.post("/{path_id}/progress")
async def update_week_progress(
    path_id: str,
    progress: WeekProgressUpdate,
    user: dict = Depends(get_current_user)
):
    """Update progress for a specific week"""
    path = await db.learning_paths.find_one(
        {"id": path_id, "user_id": user["id"]},
        {"_id": 0}
    )
    
    if not path:
        raise HTTPException(status_code=404, detail="Learning path not found")
    
    week_progress = path.get("week_progress", {})
    week_progress[str(progress.week)] = {
        "completed": progress.completed,
        "notes": progress.notes,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    completed_weeks = sum(1 for w in week_progress.values() if w.get("completed"))
    total_weeks = len(path.get("path_data", {}).get("weeks", []))
    overall_progress = round((completed_weeks / total_weeks) * 100) if total_weeks > 0 else 0
    
    await db.learning_paths.update_one(
        {"id": path_id},
        {"$set": {
            "week_progress": week_progress,
            "overall_progress": overall_progress,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {
        "message": f"Week {progress.week} progress updated",
        "week_progress": week_progress,
        "overall_progress": overall_progress
    }


@router.get("/{path_id}/progress")
async def get_path_progress(path_id: str, user: dict = Depends(get_current_user)):
    """Get progress for a learning path"""
    path = await db.learning_paths.find_one(
        {"id": path_id, "user_id": user["id"]},
        {"_id": 0}
    )
    
    if not path:
        raise HTTPException(status_code=404, detail="Learning path not found")
    
    return {
        "week_progress": path.get("week_progress", {}),
        "course_progress": path.get("course_progress", {}),
        "overall_progress": path.get("overall_progress", 0)
    }


@router.post("/{path_id}/course-progress")
async def update_course_progress(
    path_id: str,
    progress: CourseProgressUpdate,
    user: dict = Depends(get_current_user)
):
    """Update progress for a specific course within a week"""
    path = await db.learning_paths.find_one(
        {"id": path_id, "user_id": user["id"]},
        {"_id": 0}
    )
    
    if not path:
        raise HTTPException(status_code=404, detail="Learning path not found")
    
    course_progress = path.get("course_progress", {})
    week_key = str(progress.week)
    
    if week_key not in course_progress:
        course_progress[week_key] = {}
    
    course_progress[week_key][str(progress.course_index)] = {
        "completed": progress.completed,
        "notes": progress.notes,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    weeks_data = path.get("path_data", {}).get("weeks", [])
    total_courses = sum(len(w.get("courses", [])) for w in weeks_data)
    completed_courses = sum(
        1 for week_courses in course_progress.values()
        for course in week_courses.values()
        if course.get("completed")
    )
    overall_progress = round((completed_courses / total_courses) * 100) if total_courses > 0 else 0
    
    await db.learning_paths.update_one(
        {"id": path_id},
        {"$set": {
            "course_progress": course_progress,
            "overall_progress": overall_progress,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {
        "message": f"Course progress updated",
        "course_progress": course_progress,
        "overall_progress": overall_progress
    }


@router.get("/{path_id}/course-progress")
async def get_course_progress(path_id: str, user: dict = Depends(get_current_user)):
    """Get all course progress for a learning path"""
    path = await db.learning_paths.find_one(
        {"id": path_id, "user_id": user["id"]},
        {"_id": 0}
    )
    
    if not path:
        raise HTTPException(status_code=404, detail="Learning path not found")
    
    course_progress = path.get("course_progress", {})
    week_progress = path.get("week_progress", {})
    
    weeks_data = path.get("path_data", {}).get("weeks", [])
    total_courses = sum(len(w.get("courses", [])) for w in weeks_data)
    completed_courses = sum(
        1 for week_courses in course_progress.values()
        for course in week_courses.values()
        if course.get("completed")
    )
    
    return {
        "course_progress": course_progress,
        "week_progress": week_progress,
        "overall_progress": path.get("overall_progress", 0),
        "stats": {
            "total_courses": total_courses,
            "completed_courses": completed_courses,
            "completion_percentage": round((completed_courses / total_courses) * 100) if total_courses > 0 else 0
        }
    }
